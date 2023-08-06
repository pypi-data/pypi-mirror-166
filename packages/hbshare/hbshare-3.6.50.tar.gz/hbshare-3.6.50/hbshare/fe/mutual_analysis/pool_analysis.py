import datetime
import pandas as pd
import numpy as np
from hbshare.fe.XZ import db_engine
from hbshare.fe.XZ import functionality
from hbshare.fe.mutual_analysis import  jj_picturing as jjpic
# from ortools.linear_solver import  pywraplp
from PIL import Image
from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


util=functionality.Untils()
hbdb=db_engine.HBDB()
localdb=db_engine.PrvFunDB().engine
plot=functionality.Plot(1000,600)



def get_core_pool():

    sql="select jjdm,jjjc from st_fund.t_st_gm_jjxx where hxbz='1' and cpfl='2' "
    pool=hbdb.db2df(sql,db='funduser')

    return  pool

def pool_his_fromexcel2db():

    df_in=pd.read_excel(r"C:\Users\xuhuai.zhe\Downloads\公募股多核心池20220429.xlsx",
                     sheet_name='核心池-主动股多').sort_values('调入时间').drop(['备注','研究员'],axis=1)
    df_out=pd.read_excel(r"C:\Users\xuhuai.zhe\Downloads\公募股多核心池20220429.xlsx",
                     sheet_name='出池').sort_values('调出时间').drop('备注',axis=1)
    df_in['调入时间']=df_in['调入时间'].astype(str).str[0:7].str.replace('-', '')
    df_in['调出时间']=''
    df_out['调出时间'] = df_out['调出时间'].astype(str).str[0:7].str.replace('-', '')
    df_out['调入时间'] = df_out['调入时间'].astype(str).str[0:7].str.replace('-', '')
    df=pd.concat([df_in,df_out],axis=0).sort_values('调入时间')
    df.reset_index(drop=True,inplace=True)
    df['基金代码']=df['基金代码'].astype(str).str[0:6]

    poo_his=pd.DataFrame()
    for year in ['2020','2021','2022']:
        for month in ['03','06','09','12']:
            tempdf=df[(df['调入时间']<=year+month)&((df['调出时间']=='')|(df['调出时间']>year+month))][['基金代码','基金名称','基金经理']]
            tempdf['asofdate']=year+month
            poo_his=pd.concat([poo_his,tempdf],axis=0)

    poo_his.reset_index(drop=True,inplace=True)
    poo_his=poo_his[poo_his['asofdate']<=str(datetime.datetime.now())[0:7].replace('-','')]
    sql='delete from core_pool_history'
    localdb.execute(sql)
    poo_his.to_sql('core_pool_history',con=localdb,if_exists='append',index=False)

def get_core_poolfromlocaldb(asofdate=None):

    if(asofdate is not None):
        sql="select * from core_pool_history where asofdate='{}'".format(asofdate)
    else:
        sql="select * from core_pool_history"
    pool=pd.read_sql(sql,con=localdb).rename(columns={'基金代码':'jjdm',
                                                      '基金名称':'jjjc',
                                                      '基金经理':'manager'})

    return pool

def get_jj_picture(jjdm_con=None,if_inverse=False):

    if(if_inverse):
        inornot='not in '
    else:
        inornot='in'

    value_df_list = [pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(),
                     pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()]

    industry_df_list = [pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame()]


    stock_df_list = [pd.DataFrame(), pd.DataFrame()]


    if(jjdm_con is not None):
        value_p, value_p_hbs, value_sp, value_sp_hbs, size_p, size_p_hbs, size_sp, size_sp_hbs, \
        industry_p, industry_sp, theme_p, theme_sp, industry_detail_df_list, stock_p, stock_tp=\
            jjpic.get_pic_from_localdb("jjdm {1} ({0})".format(jjdm_con,inornot), if_percentage=False)
    else:
        value_p, value_p_hbs, value_sp, value_sp_hbs, size_p, size_p_hbs, size_sp, size_sp_hbs, \
        industry_p, industry_sp, theme_p, theme_sp, industry_detail_df_list, stock_p, stock_tp=\
            jjpic.get_pic_from_localdb("1=1", if_percentage=False)

    value_df_list[0]=value_p
    value_df_list[1] = value_p_hbs
    value_df_list[2] = value_sp
    value_df_list[3] = value_sp_hbs
    value_df_list[4] = size_p
    value_df_list[5] = size_p_hbs
    value_df_list[6] = size_sp
    value_df_list[7] = size_sp_hbs

    industry_df_list[0]=industry_p
    industry_df_list[1] = industry_sp
    industry_df_list[2] = theme_p
    industry_df_list[3] = theme_sp
    for j in range(len(industry_detail_df_list)):
        # industry_detail_df_list[j]['jjdm'] = jjdm
        industry_detail_df_list[j]['industry_level'] = j + 1
        industry_df_list[4] = pd.concat([industry_df_list[4], industry_detail_df_list[j]], axis=0)

    stock_df_list[0]=stock_p
    stock_df_list[1] = stock_tp

    return value_df_list,industry_df_list,stock_df_list

def features_calculation(value_df_list, industry_df_list, stock_df_list):

    industry_style=industry_df_list[0].groupby('一级行业类型').count()['jjdm'].to_frame('num')
    industry_style=industry_style/len(industry_df_list[0])*100
    industry_style_dis = (industry_df_list[4].groupby(['industry_level', '行业名称']).mean() * 100).reset_index()[
        ['industry_level', '行业名称', '占持仓比例(时序均值)', '占持仓比例(时序均值)排名']]
    #take industry with weight more than 5%
    industry_style_dis=industry_style_dis[industry_style_dis['占持仓比例(时序均值)'] >= (5/industry_style_dis['industry_level'])]
    class_list=[]
    for i in range(3):
        class_list.append(industry_style_dis[industry_style_dis['industry_level'] == (i + 1)] \
            .sort_values('占持仓比例(时序均值)', ascending=False).set_index('行业名称').drop('industry_level', axis=1))

    style_type_dis = value_df_list[1].groupby('风格类型').count()['jjdm'].to_frame('风格类型') \
                     / len(value_df_list[1]) * 100
    style_incline_dis=value_df_list[1][value_df_list[1]['风格类型']=='专注'][['风格类型','风格偏好']]
    style_incline_dis=style_incline_dis.groupby('风格偏好').count()/len(style_incline_dis)*100
    style_weight_dis=value_df_list[1][['成长绝对暴露(持仓)',
       '价值绝对暴露(持仓)','成长暴露排名(持仓)', '价值暴露排名(持仓)']].mean(axis=0).to_frame('成长价值权重分布')*100

    style_type_dis2 = value_df_list[5].groupby('规模风格类型').count()['jjdm'].to_frame('规模风格类型') \
                      / len(value_df_list[5]) * 100
    style_incline_dis2=value_df_list[5][value_df_list[5]['规模风格类型']=='专注'][['规模风格类型','规模偏好']]
    style_incline_dis2=style_incline_dis2.groupby('规模偏好').count()/len(style_incline_dis2)*100
    style_weight_dis2=value_df_list[5][['大盘绝对暴露(持仓)',
       '中盘绝对暴露(持仓)', '小盘绝对暴露(持仓)','大盘暴露排名(持仓)',
       '中盘暴露排名(持仓)', '小盘暴露排名(持仓)']].mean(axis=0).to_frame('成长价值权重分布')*100

    stock_style_a=stock_df_list[0].groupby('个股风格A').count()['jjdm'].to_frame('个股风格类型A')\
                  /len(stock_df_list[0])*100
    stock_style_b = stock_df_list[0].groupby('个股风格B').count()['jjdm'].to_frame('个股风格类型B')\
                    /len(stock_df_list[0])*100
    stock_left = stock_df_list[1].groupby('左侧标签').count()['jjdm'].to_frame('左侧类型分布')\
                 / len(stock_df_list[1]) * 100
    stock_df_list[1].loc[stock_df_list[1]['新股次新股偏好'] == '', '新股次新股偏好'] = '无'
    stock_new=stock_df_list[1].groupby('新股次新股偏好').count()['jjdm'].to_frame('新股偏好分布')\
              / len(stock_df_list[1]) * 100
    stock_fin2=(stock_df_list[1][['平均持有时间(出持仓前)_rank', '左侧概率(出重仓前,半年线)_rank','新股概率(出持仓前)_rank'
        ,'出重仓前平均收益率_rank',
       '出全仓前平均收益率_rank',]]).mean(axis=0)
    stock_fin=(stock_df_list[0][['PE_rank', 'PB_rank',
       'ROE_rank', '股息率_rank','平均仓位']]).mean(axis=0)
    stock_fin=pd.concat([stock_fin,stock_fin2],axis=0).to_frame('个股持仓特征')
    stock_fin.index = [x.replace('rank', '排名') for x in stock_fin.index]
    stock_fin=stock_fin*100


    return industry_style,class_list,style_type_dis,style_incline_dis,\
           style_weight_dis,style_type_dis2,style_incline_dis2,style_weight_dis2,\
           stock_style_a,stock_style_b,stock_left,stock_new,stock_fin

def pool_industry(industry_style,class_list,save_local_file=False):

    if(len(industry_style.columns)>1):
        plot.plotly_jjpic_bar(industry_style, '行业风格分布',save_local_file=save_local_file)
    else:
        plot.plotly_pie(industry_style, '行业风格分布',save_local_file=save_local_file)

    for i in range(3):

        class_list[i].columns=[x.replace('占持仓比例(时序均值)','') for x in class_list[i].columns]

        plot.plotly_jjpic_bar(class_list[i][class_list[i].columns[0:int(len(class_list[i].columns)/2)]].fillna(0),"{}级行业持仓分布_占持仓比例".format(i+1),save_local_file=save_local_file)
        plot.plotly_jjpic_bar(class_list[i][class_list[i].columns[int(len(class_list[i].columns)/2):]].fillna(0), "{}级行业持仓分布_占持仓比例排名".format(i + 1),save_local_file=save_local_file)

def pool_style(style_type_dis,style_incline_dis,
               style_weight_dis,style_type_dis2,style_incline_dis2,style_weight_dis2,save_local_file=False):

    if(len(style_type_dis.columns)>1):
        plot.plotly_jjpic_bar(style_type_dis, '风格类型分布',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_incline_dis, '风格专注型分布',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_weight_dis.iloc[0:2], '成长价值持仓占比',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_weight_dis.iloc[2:], '成长价值持仓占比排名',save_local_file=save_local_file)

        plot.plotly_jjpic_bar(style_type_dis2, '规模风格类型分布',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_incline_dis2, '专注型风格分布',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_weight_dis2.iloc[0:3], '大中小盘持仓占比',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(style_weight_dis2.iloc[3:], '大中小盘持仓占比排名',save_local_file=save_local_file)
    else:
        plot.plotly_pie(style_type_dis,'风格类型分布',save_local_file=save_local_file)
        plot.plotly_pie(style_incline_dis,'风格专注型分布',save_local_file=save_local_file)
        plot.plotly_pie(style_weight_dis.iloc[0:2],'成长价值持仓占比',save_local_file=save_local_file)
        plot.plotly_pie(style_weight_dis.iloc[2:],'成长价值持仓占比排名',save_local_file=save_local_file)


        plot.plotly_pie(style_type_dis2,'规模风格类型分布',save_local_file=save_local_file)
        plot.plotly_pie(style_incline_dis2,'专注型风格分布',save_local_file=save_local_file)
        plot.plotly_pie(style_weight_dis2.iloc[0:3],'大中小盘持仓占比',save_local_file=save_local_file)
        plot.plotly_pie(style_weight_dis2.iloc[3:],'大中小盘持仓占比排名',save_local_file=save_local_file)

def pool_other(stock_style_a,stock_style_b,stock_left,stock_new,stock_fin,save_local_file=False):

    if(len(stock_style_a.columns)>1):
        plot.plotly_jjpic_bar(stock_style_a,'个股风格类型分布A',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(stock_style_b,'个股风格类型分布B',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(stock_left ,'左侧类型分布',save_local_file=save_local_file)
        plot.plotly_jjpic_bar(stock_new,'新股偏好分布',save_local_file=save_local_file)
    else:
        plot.plotly_pie(stock_style_a,'个股风格类型分布A',save_local_file=save_local_file)
        plot.plotly_pie(stock_style_b,'个股风格类型分布B',save_local_file=save_local_file)
        plot.plotly_pie(stock_left ,'左侧类型分布',save_local_file=save_local_file)
        plot.plotly_pie(stock_new,'新股偏好分布',save_local_file=save_local_file)
    plot.plotly_jjpic_bar(stock_fin,'个股持仓特征',save_local_file=save_local_file)

def pool_picturing_engine(industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin,save_local_file=False):

    pool_industry(industry_style,class_list,save_local_file)
    pool_style(style_type_dis,style_incline_dis,
               style_weight_dis,style_type_dis2,style_incline_dis2,style_weight_dis2,save_local_file)
    pool_other(stock_style_a,stock_style_b,stock_left,stock_new,stock_fin,save_local_file)


def get_potient_pool(jjdm_con):

    value_df_list,industry_df_list,stock_df_list=get_jj_picture(jjdm_con,if_inverse=True)

    return value_df_list,industry_df_list,stock_df_list

# def optimizer_allocation(value_df_list,industry_df_list, stock_df_list,pool,industry_style, class_list, style_type_dis, style_incline_dis, \
#     style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
#     stock_style_a, stock_style_b, stock_left, stock_new, stock_fin):
#
#     turnover_num=20
#     style_order = ['专注', '博弈', '轮动', '配置']
#
#     jjdm_con=util.list_sql_condition(pool['jjdm'].unique().tolist())
#     value_df_list_pot, industry_df_list_pot, stock_df_list_pot=get_potient_pool(jjdm_con)
#
#     potient_jj_list=industry_df_list_pot[0]['jjdm'].tolist()
#
#
#     sql="select jjdm,zxppm1nyj from st_fund.t_st_gm_nxpblpm where jjdm not in ({0}) and tjnf>={1} ".format(jjdm_con,2022-1)
#     sharpr=-1*hbdb.db2df(sql,db='funduser').set_index('jjdm').loc[potient_jj_list]['zxppm1nyj'].astype(float).values
#
#     solver=pywraplp.Solver.CreateSolver('SCIP')
#     x=[]
#     for i in potient_jj_list:
#         x.append(solver.IntVar(0,1,"x_{}".format(i)))
#
#
#     input_industry_style_cons_ub=np.array([0.3,0.3,0.3,0.3])
#     input_industry_style_cons_lb = np.array([0, 0, 0, 0])
#     input_industry_style_cons_ub=input_industry_style_cons_ub*(turnover_num+len(pool))\
#                                  -industry_style.loc[style_order]['num'].values/100*len(pool)
#     input_industry_style_cons_lb=input_industry_style_cons_lb*(turnover_num+len(pool))\
#                                  -industry_style.loc[style_order]['num'].values/100*len(pool)
#
#
#     cons_dict=dict()
#     #set constrains bound
#     cons_dict['turnover_num']=solver.Constraint(turnover_num, turnover_num)
#     for m in range(4):
#         cons_dict['industry_style_{}'.format(style_order[m])] = solver.Constraint(input_industry_style_cons_lb[m], input_industry_style_cons_ub[m])
#     industry_style_cofe=pd.get_dummies(industry_df_list_pot[0].set_index('jjdm')
#                                        .loc[potient_jj_list]['一级行业类型'])[style_order].astype(float).values
#
#     for i in range(len(x)):
#         cons_dict['turnover_num'].SetCoefficient(x[i], 1)
#         for m in range(4):
#             cons_dict['industry_style_{}'.format(style_order[m])].SetCoefficient(x[i], industry_style_cofe[i,m])
#
#     objective = solver.Objective()
#     for i in range(len(x)):
#         objective.SetCoefficient(x[i], sharpr[i])
#     objective.SetMaximization()
#
#     result_status = solver.Solve()
#     print(result_status == pywraplp.Solver.OPTIMAL)
#     print(solver.Objective().Value())
#
#     adding_list=[]
#     for i in range(len(x)):
#         if(x[i].solution_value()==1):
#             adding_list.append(str(x[i]).split('_')[-1])
#
#     new_jjdm_list=pool['jjdm'].unique().tolist()+adding_list
#
#     return  new_jjdm_list

def pool_picturing_opt():

    pool = get_core_pool()
    pool=get_core_poolfromlocaldb('202203')
    #picture the original pool
    jjdm_con = util.list_sql_condition(pool['jjdm'].unique().tolist())
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(jjdm_con)
    industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin\
        =features_calculation(value_df_list, industry_df_list, stock_df_list)

    pool_picturing_engine(industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin)

    # new_jjdm_list=optimizer_allocation(value_df_list, industry_df_list, stock_df_list,pool,industry_style, class_list, style_type_dis, style_incline_dis, \
    # style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    # stock_style_a, stock_style_b, stock_left, stock_new, stock_fin)

    #pic the new pool
    #new_jjdm_con=util.list_sql_condition(new_jjdm_list)
    new_jjdm_con="'001705','001410','001856','450009','166006'," \
                 "'001975','003095','005241','161606','005827','002708'," \
                 "'005968','005267','519133','002340','000577','163415'," \
                 "'000729','000756','000991','001208','001487','001532'," \
                 "'001718','001877','002258','002871','002943','004350'," \
                 "'004784','005161','005457','005669','005825','006257','007130','519702'"
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(new_jjdm_con)
    industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new\
        =features_calculation(value_df_list, industry_df_list, stock_df_list)

    industry_style_new=pd.merge(industry_style,industry_style_new,how='outer',left_index=True,right_index=True)
    industry_style_new.columns=['旧池','新池']

    for i in range(3):
        class_list_new[i]=pd.merge(class_list[i],class_list_new[i],how='outer',
                                   left_index=True,right_index=True)
        class_list_new[i].columns=[x.replace('x','旧') for x in class_list_new[i].columns]
        class_list_new[i].columns = [x.replace('y', '新') for x in class_list_new[i].columns]
        class_list_new[i]=class_list_new[i][class_list_new[i].columns.sort_values()]


    style_type_dis_new=pd.merge(style_type_dis,style_type_dis_new,how='outer',
                                left_index=True,right_index=True).fillna(0)
    style_type_dis_new.columns=['旧池','新池']

    style_incline_dis_new=pd.merge(style_incline_dis,style_incline_dis_new,how='outer',
                                   left_index=True,right_index=True).fillna(0)
    style_incline_dis_new.columns=['旧池','新池']

    style_weight_dis_new=pd.merge(style_weight_dis,style_weight_dis_new,how='outer',
                                  left_index=True,right_index=True).fillna(0)
    style_weight_dis_new.columns=['旧池','新池']

    style_type_dis2_new=pd.merge(style_type_dis2,style_type_dis2_new,how='outer',
                                 left_index=True,right_index=True).fillna(0)
    style_type_dis2_new.columns=['旧池','新池']

    style_incline_dis2_new=pd.merge(style_incline_dis2,style_incline_dis2_new,
                                    how='outer',left_index=True,right_index=True).fillna(0)
    style_incline_dis2_new.columns=['旧池','新池']

    style_weight_dis2_new=pd.merge(style_weight_dis2,style_weight_dis2_new,
                                   how='outer',left_index=True,right_index=True).fillna(0)
    style_weight_dis2_new.columns=['旧池','新池']

    stock_style_a_new=pd.merge(stock_style_a,stock_style_a_new,
                               how='outer',left_index=True,right_index=True).fillna(0)
    stock_style_a_new.columns=['旧池','新池']

    stock_style_b_new=pd.merge(stock_style_b,stock_style_b_new,how='outer'
                               ,left_index=True,right_index=True).fillna(0)
    stock_style_b_new.columns=['旧池','新池']

    stock_left_new=pd.merge(stock_left,stock_left_new,how='outer'
                            ,left_index=True,right_index=True).fillna(0)
    stock_left_new.columns=['旧池','新池']

    stock_new_new=pd.merge(stock_new,stock_new_new,how='outer',
                           left_index=True,right_index=True).fillna(0)
    stock_new_new.columns=['旧池','新池']

    stock_fin_new=pd.merge(stock_fin,stock_fin_new,how='outer',
                           left_index=True,right_index=True).fillna(0)
    stock_fin_new.columns=['旧池','新池']


    pool_picturing_engine(industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new)

def pool_picturing(pool=None,returndata=False,save_local_file=False):

    if(pool is None):
        pool=get_core_poolfromlocaldb('202203')

    #picture the original pool
    jjdm_con = util.list_sql_condition(pool['jjdm'].unique().tolist())
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(jjdm_con)
    industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin\
        =features_calculation(value_df_list, industry_df_list, stock_df_list)

    plot.plotly_table(pool, 1000, '池列表', save_local_file=True)
    pool_picturing_engine(industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin,save_local_file)

    if(returndata):
        return industry_style, class_list, style_type_dis, style_incline_dis, \
        style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
        stock_style_a, stock_style_b, stock_left, stock_new, stock_fin

def pool_comparision(date1,date2):

    pool=get_core_poolfromlocaldb(date1)

    #picture the original pool
    jjdm_con = util.list_sql_condition(pool['jjdm'].unique().tolist())
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(jjdm_con)
    industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin\
        =features_calculation(value_df_list, industry_df_list, stock_df_list)

    # picture the new pool
    new_pool = get_core_poolfromlocaldb(date2)
    new_jjdm_con=util.list_sql_condition(new_pool['jjdm'].unique().tolist())
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(new_jjdm_con)
    industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new\
        =features_calculation(value_df_list, industry_df_list, stock_df_list)

    #get the return for in and out jj for the next 1 quarter and year
    in_jjdm=list(set(new_pool['jjdm'].unique().tolist()).difference(set(pool['jjdm'].unique().tolist())))
    out_jjdm=list(set(pool['jjdm'].unique().tolist()).difference(set(new_pool['jjdm'].unique().tolist())))
    start_date=util._shift_date(date2+'31')
    in_ret=pd.DataFrame(data=in_jjdm,columns=['jjdm'])
    out_ret = pd.DataFrame(data=out_jjdm, columns=['jjdm'])

    end_date_y=util._shift_date(str(int(date2[0:4])+1)+date2[4:]+'30')
    if(date2[4:]!='12'):
        end_date_q=util._shift_date(date2[0:4]+("0"+str(int(date2[4:])+3))[-2:]+'30')
    else:
        end_date_q=util._shift_date(str(int(date2[0:4])+1)+'0331')

    if(end_date_y>=datetime.datetime.today().strftime('%Y%m%d')[0:6]):
        end_date_y=util._shift_date(str(hbdb.db2df("select max(jzrq) as jzrq from st_fund.t_st_gm_rhb", db='funduser')['jzrq'][0]))
    sql = "select jjdm,jzrq,fqdwjz from st_fund.t_st_gm_rhb where (jzrq='{0}' or  jzrq='{1}') and jjdm in({2})" \
        .format(end_date_y, start_date, util.list_sql_condition(in_jjdm+out_jjdm))
    gap_return_y = hbdb.db2df(sql, db='funduser')
    gap_return_y=pd.merge(gap_return_y[gap_return_y['jzrq']==int(start_date)]
                          ,gap_return_y[gap_return_y['jzrq']==int(end_date_y)],on='jjdm')
    gap_return_y['ret']=gap_return_y['fqdwjz_y']/gap_return_y['fqdwjz_x']-1


    if(end_date_q>=datetime.datetime.today().strftime('%Y%m%d')[0:6]):
        end_date_q = util._shift_date(str(hbdb.db2df("select max(jzrq) as jzrq from st_fund.t_st_gm_rhb", db='funduser')['jzrq'][0]))
    sql = "select jjdm,jzrq,fqdwjz from st_fund.t_st_gm_rhb where (jzrq='{0}' or  jzrq='{1}') and jjdm in({2})" \
        .format(end_date_q, start_date, util.list_sql_condition(in_jjdm+out_jjdm))
    gap_return_q = hbdb.db2df(sql, db='funduser')
    gap_return_q=pd.merge(gap_return_q[gap_return_q['jzrq']==int(start_date)]
                          ,gap_return_q[gap_return_q['jzrq']==int(end_date_q)],on='jjdm')
    gap_return_q['ret']=gap_return_q['fqdwjz_y']/gap_return_q['fqdwjz_x']-1

    in_ret=pd.merge(in_ret,gap_return_q[['jjdm','ret']],how='left',on='jjdm')\
        .rename(columns={'ret':'ret_q'})
    in_ret=pd.merge(in_ret,gap_return_y[['jjdm','ret']],how='left',on='jjdm')\
        .rename(columns={'ret':'ret_y'})
    in_ret=in_ret.mean()
    out_ret=pd.merge(out_ret,gap_return_q[['jjdm','ret']],how='left',on='jjdm')\
        .rename(columns={'ret':'ret_q'})
    out_ret=pd.merge(out_ret,gap_return_y[['jjdm','ret']],how='left',on='jjdm')\
        .rename(columns={'ret':'ret_y'})
    out_ret = out_ret.mean()

    #join the two pool picture
    industry_style_new=pd.merge(industry_style,industry_style_new,how='outer',left_index=True,right_index=True)
    industry_style_new.columns=['旧池','新池']

    for i in range(3):
        class_list_new[i]=pd.merge(class_list[i],class_list_new[i],how='outer',
                                   left_index=True,right_index=True)
        class_list_new[i].columns=[x.replace('x','旧') for x in class_list_new[i].columns]
        class_list_new[i].columns = [x.replace('y', '新') for x in class_list_new[i].columns]
        class_list_new[i]=class_list_new[i][class_list_new[i].columns.sort_values()]


    style_type_dis_new=pd.merge(style_type_dis,style_type_dis_new,how='outer',
                                left_index=True,right_index=True).fillna(0)
    style_type_dis_new.columns=['旧池','新池']

    style_incline_dis_new=pd.merge(style_incline_dis,style_incline_dis_new,how='outer',
                                   left_index=True,right_index=True).fillna(0)
    style_incline_dis_new.columns=['旧池','新池']

    style_weight_dis_new=pd.merge(style_weight_dis,style_weight_dis_new,how='outer',
                                  left_index=True,right_index=True).fillna(0)
    style_weight_dis_new.columns=['旧池','新池']

    style_type_dis2_new=pd.merge(style_type_dis2,style_type_dis2_new,how='outer',
                                 left_index=True,right_index=True).fillna(0)
    style_type_dis2_new.columns=['旧池','新池']

    style_incline_dis2_new=pd.merge(style_incline_dis2,style_incline_dis2_new,
                                    how='outer',left_index=True,right_index=True).fillna(0)
    style_incline_dis2_new.columns=['旧池','新池']

    style_weight_dis2_new=pd.merge(style_weight_dis2,style_weight_dis2_new,
                                   how='outer',left_index=True,right_index=True).fillna(0)
    style_weight_dis2_new.columns=['旧池','新池']

    stock_style_a_new=pd.merge(stock_style_a,stock_style_a_new,
                               how='outer',left_index=True,right_index=True).fillna(0)
    stock_style_a_new.columns=['旧池','新池']

    stock_style_b_new=pd.merge(stock_style_b,stock_style_b_new,how='outer'
                               ,left_index=True,right_index=True).fillna(0)
    stock_style_b_new.columns=['旧池','新池']

    stock_left_new=pd.merge(stock_left,stock_left_new,how='outer'
                            ,left_index=True,right_index=True).fillna(0)
    stock_left_new.columns=['旧池','新池']

    stock_new_new=pd.merge(stock_new,stock_new_new,how='outer',
                           left_index=True,right_index=True).fillna(0)
    stock_new_new.columns=['旧池','新池']

    stock_fin_new=pd.merge(stock_fin,stock_fin_new,how='outer',
                           left_index=True,right_index=True).fillna(0)
    stock_fin_new.columns=['旧池','新池']



    pool_picturing_engine(industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new)


    return  in_ret,out_ret

def pic_pool_history(returndata=False,save_local_file=False):

    pool_his=get_core_poolfromlocaldb()

    jjdm_con = util.list_sql_condition(pool_his['jjdm'].unique().tolist())
    value_df_list, industry_df_list, stock_df_list = get_jj_picture(jjdm_con)

    industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new=pd.DataFrame(),[pd.DataFrame(),pd.DataFrame(),pd.DataFrame()],\
                                                                                       pd.DataFrame(),pd.DataFrame(),\
                                                                                       pd.DataFrame(),pd.DataFrame(),pd.DataFrame(),\
                                                                                       pd.DataFrame(),pd.DataFrame(),pd.DataFrame(),\
                                                                                       pd.DataFrame(),pd.DataFrame(),pd.DataFrame()

    new_col_list = []
    for asofdate in pool_his['asofdate'].unique():
        jjdm_df=pool_his[pool_his['asofdate']==asofdate]['jjdm'].to_frame('jjdm_x')
        new_col_list.append(str(asofdate))
        value_df_list_temp=[]
        industry_df_list_temp=[]
        stock_df_list_temp=[]
        for i in range(len(value_df_list)):
            value_df_list_temp.append(pd.merge(jjdm_df,value_df_list[i],
                                           how='left',left_on='jjdm_x',right_on='jjdm').drop('jjdm_x',axis=1))
        for i in range(len(industry_df_list)):
            industry_df_list_temp.append(pd.merge(jjdm_df,industry_df_list[i],
                                           how='left',left_on='jjdm_x',right_on='jjdm').drop('jjdm_x',axis=1))
        for i in range(len(stock_df_list)):
            stock_df_list_temp.append(pd.merge(jjdm_df,stock_df_list[i],
                                           how='left',left_on='jjdm_x',right_on='jjdm').drop('jjdm_x',axis=1))

        industry_style, class_list, style_type_dis, style_incline_dis, \
        style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
        stock_style_a, stock_style_b, stock_left, stock_new, stock_fin\
            =features_calculation(value_df_list_temp, industry_df_list_temp, stock_df_list_temp)

        industry_style_new = pd.merge(industry_style_new,industry_style,  how='outer', left_index=True,
                                      right_index=True)


        for i in range(3):

            class_list[i].columns = [x+"_"+str(asofdate) for x in class_list[i].columns]
            class_list_new[i] = pd.merge( class_list_new[i],class_list[i], how='outer',
                                         left_index=True, right_index=True)
            class_list_new[i] = class_list_new[i][class_list_new[i].columns.sort_values()]

        style_type_dis_new = pd.merge( style_type_dis_new,style_type_dis, how='outer',
                                      left_index=True, right_index=True).fillna(0)


        style_incline_dis_new = pd.merge(style_incline_dis_new,style_incline_dis, how='outer',
                                         left_index=True, right_index=True).fillna(0)


        style_weight_dis_new = pd.merge( style_weight_dis_new,style_weight_dis, how='outer',
                                        left_index=True, right_index=True).fillna(0)


        style_type_dis2_new = pd.merge(style_type_dis2_new,style_type_dis2,  how='outer',
                                       left_index=True, right_index=True).fillna(0)


        style_incline_dis2_new = pd.merge( style_incline_dis2_new,style_incline_dis2,
                                          how='outer', left_index=True, right_index=True).fillna(0)


        style_weight_dis2_new = pd.merge(style_weight_dis2_new,style_weight_dis2,
                                         how='outer', left_index=True, right_index=True).fillna(0)


        stock_style_a_new = pd.merge(stock_style_a_new,stock_style_a,
                                     how='outer', left_index=True, right_index=True).fillna(0)


        stock_style_b_new = pd.merge( stock_style_b_new,stock_style_b, how='outer'
                                     , left_index=True, right_index=True).fillna(0)


        stock_left_new = pd.merge( stock_left_new,stock_left, how='outer'
                                  , left_index=True, right_index=True).fillna(0)


        stock_new_new = pd.merge(stock_new_new, stock_new, how='outer',
                                 left_index=True, right_index=True).fillna(0)


        stock_fin_new = pd.merge(stock_fin_new,stock_fin,  how='outer',
                                 left_index=True, right_index=True).fillna(0)


    industry_style_new.columns = new_col_list
    style_type_dis_new.columns = new_col_list
    style_incline_dis_new.columns = new_col_list
    style_weight_dis_new.columns = new_col_list
    style_type_dis2_new.columns = new_col_list
    style_incline_dis2_new.columns =new_col_list
    style_weight_dis2_new.columns = new_col_list
    stock_style_a_new.columns = new_col_list
    stock_style_b_new.columns = new_col_list
    stock_left_new.columns = new_col_list
    stock_new_new.columns = new_col_list
    stock_fin_new.columns = new_col_list


    pool_picturing_engine(industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new,save_local_file=save_local_file)


    if(returndata):
        return industry_style_new, class_list_new, style_type_dis_new, style_incline_dis_new, \
    style_weight_dis_new, style_type_dis2_new, style_incline_dis2_new, style_weight_dis2_new, \
    stock_style_a_new, stock_style_b_new, stock_left_new, stock_new_new, stock_fin_new

def get_pool_history_report():
    
    def style_analysis(df, additional_str1='',additional_str2=''):
        style_style = "从基金{0}来看，其中:".format(additional_str1)
        for style in df.index:
            style_style += "{0}，".format(str(style)+additional_str2)
            style_style+="历史最高为{0}，历史最低为{1}。最新一期为{2}，最早一期为{3}"\
                .format((str(round(df.loc[style].max(), 2)) + "%,"),(str(round(df.loc[style].min(), 2)) + "%,"),
                        (str(round(df.loc[style].iloc[-1], 2)) + "%,"),(str(round(df.loc[style].iloc[0], 2)) + "%,"))

        return style_style

    pic_name_list = ['行业风格分布.png', '1级行业持仓分布_占持仓比例.png',
                     '1级行业持仓分布_占持仓比例排名.png',
                     '2级行业持仓分布_占持仓比例.png',
                     '2级行业持仓分布_占持仓比例排名.png',
                     '3级行业持仓分布_占持仓比例.png',
                     '3级行业持仓分布_占持仓比例排名.png',
                     '风格类型分布.png',
                     '风格专注型分布.png',
                     '成长价值持仓占比.png',
                     '成长价值持仓占比排名.png',
                     '规模风格类型分布.png',
                     '专注型风格分布.png',
                     '大中小盘持仓占比.png',
                     '大中小盘持仓占比排名.png',
                     '个股风格类型分布A.png',
                     '个股风格类型分布B.png',
                     '左侧类型分布.png',
                     '新股偏好分布.png',
                     '个股持仓特征.png',
                     ]

    industry_style, class_list, style_type_dis, style_incline_dis, \
    style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
    stock_style_a, stock_style_b, stock_left, stock_new, stock_fin =\
        pic_pool_history(returndata=True,save_local_file=True)


    time_length=int(len(class_list[0].columns)/2)
    report_date=class_list[0].columns[0:time_length]
    report_date_rank=class_list[0].columns[time_length:]

    data_list=[industry_style , class_list[0][report_date],class_list[0][report_date_rank],
               class_list[1][report_date],class_list[1][report_date_rank],
               class_list[2][report_date],class_list[2][report_date_rank], style_type_dis , style_incline_dis , \
    style_weight_dis[0:2] ,style_weight_dis.iloc[2:], style_type_dis2 , style_incline_dis2 , style_weight_dis2[0:3] ,style_weight_dis2[3:], \
    stock_style_a , stock_style_b , stock_left , stock_new , stock_fin]
    key_word_list=['行业配置类型','一级行业分布','一级行业分布排名','二级行业分布','二级行业分布排名','三级行业分布','三级行业分布排名',
                   '风格类型','中属于风格专注的基金','持仓中的风格占比','持仓中的风格占比排名','规模风格类型','中属于风格专注的基金','持仓中的规模风格占比','持仓中的规模风格占比排名',
                   '个股配置风格','个股策略','个股左侧特征','对次新股的偏好分布','其他个股持仓特征']
    key_word_list2 = ['型占比', '行业占比', '行业占比相对排名', '行业占比', '行业占比相对排名',
                      '行业占比', '行业占比相对排名','型占比', '型占比','占比','排名','型占比','型占比',
                      '占比','排名','型占比','型占比','型占比','型占比','']

    doc = Document()
    fold = r"E:\GitFolder\hbshare\fe\mutual_analysis"

    title_num=2
    name=''
    pic_paragraphs_list = []
    string_paragraphs_list = []
    paragraphs_count = 0

    doc.add_paragraph("")
    paragraphs_count += 1

    p = doc.add_paragraph("按照季度，对公墓核心池历史进行了统计分析，结果如下")
    p.style.font.size = Pt(10)
    p.paragraph_format.first_line_indent = p.style.font.size * 2
    string_paragraphs_list.append(paragraphs_count)
    paragraphs_count += 1

    for i in range(len(pic_name_list)):
        pic=pic_name_list[i]
        file_path = fold + "\\" + pic
        desc = style_analysis(data_list[i].fillna(0), key_word_list[i],key_word_list2[i])

        doc.add_paragraph("")
        paragraphs_count += 1

        try:
            doc.add_picture(file_path, width=Inches(6), height=Inches(4))
            pic_paragraphs_list.append(paragraphs_count)
            paragraphs_count += 1

        except Exception as e:
            pic_temp = Image.open(file_path)
            pic_temp.save(pic_temp)
            doc.add_picture(file_path, width=Inches(6), height=Inches(4))

        p = doc.add_paragraph(desc)
        p.style.font.size = Pt(10)
        p.paragraph_format.first_line_indent = p.style.font.size * 2
        string_paragraphs_list.append(paragraphs_count)
        paragraphs_count += 1

    for j in pic_paragraphs_list:
        doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for j in string_paragraphs_list:
        doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.save(r"E:\GitFolder\hbshare\fe\mutual_analysis\公募核心池历史报告.docx")
    print('')


if __name__ == '__main__':

    get_pool_history_report()
    #pool_his_fromexcel2db()
    #pic_pool_history()
    #pool_comparision('202103', '202203')

    #pool_picturing