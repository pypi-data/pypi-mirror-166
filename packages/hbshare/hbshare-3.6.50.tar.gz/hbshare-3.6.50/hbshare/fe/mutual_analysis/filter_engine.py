import datetime
import pandas as pd
import numpy as np
from hbshare.fe.XZ import db_engine
from hbshare.fe.XZ import functionality
from hbshare.fe.mutual_analysis import pool_analysis as pla




util=functionality.Untils()
hbdb=db_engine.HBDB()
localdb=db_engine.PrvFunDB().engine

jjdm_list = util.get_mutual_stock_funds(datetime.datetime.today().strftime('%Y%m%d'))
jjdm_con = util.list_sql_condition(jjdm_list)
sql = "select jjjc,jjdm from st_fund.t_st_gm_jjxx where jjdm in ({0})".format(jjdm_con)
jjdm_name_map = hbdb.db2df(sql, db='funduser')

class Equilibrium:

    @staticmethod
    def ind_equ():

        latest_asofdate=\
            pd.read_sql("select max(asofdate) as asofdate from hbs_industry_property_new ",con=localdb)['asofdate'][0]

        sql="SELECT jjdm,cen_ind_1,asofdate from hbs_industry_property_new where asofdate='{0}' order by cen_ind_1 ".format(latest_asofdate)
        ind_equ=pd.read_sql(sql,con=localdb)

        return ind_equ

    @staticmethod
    def style_equ():

        latest_asofdate = \
        pd.read_sql("select max(asofdate) as asofdate from hbs_style_property ", con=localdb)['asofdate'][0]

        sql="""SELECT jjdm,cen_lv_rank,asofdate 
        from hbs_style_property where asofdate='{0}' order by cen_lv_rank 
        """.format(latest_asofdate)
        style_equ = pd.read_sql(sql, con=localdb)

        return  style_equ

    @staticmethod
    def size_equ():

        latest_asofdate = \
        pd.read_sql("select max(asofdate) as asofdate from hbs_size_property ", con=localdb)['asofdate'][0]

        sql = """SELECT jjdm,cen_lv_rank,asofdate 
        from hbs_size_property where asofdate='{0}' order by cen_lv_rank 
        """.format(latest_asofdate)

        size_equ = pd.read_sql(sql, con=localdb)

        return  size_equ

    @staticmethod
    def pepbroe_equ():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_holding_property ", con=localdb)[
                'asofdate'][0]

        sql="""
        select jjdm,PE_rank,PB_rank,ROE_rank,股息率_rank,asofdate 
        from hbs_holding_property where asofdate='{0}'
        """.format(latest_asofdate)

        pepbroe_equ=pd.read_sql(sql,con=localdb)
        pepbroe_equ[
            ['PE_rank','PB_rank','ROE_rank','股息率_rank']
        ]=abs(pepbroe_equ[['PE_rank','PB_rank','ROE_rank','股息率_rank']]-0.5)

        pepbroe_equ=pepbroe_equ.sort_values('PE_rank')

        return pepbroe_equ

    @staticmethod
    def nav_equ():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from nav_ret_bias ", con=localdb)[
                'asofdate'][0]

        sql="""
        select jjdm,mean_rank_monthly,std_rank_monthly,mean_rank_weekly,std_rank_weekly,asofdate
        from nav_ret_bias where asofdate='{0}' order by mean_rank_monthly
        """.format(latest_asofdate)

        nav_equ=pd.read_sql(sql,con=localdb)

        return  nav_equ

    def get_equilibrium(self,threshield,show_num=300):

        # method_1
        ind_equ = self.ind_equ()
        style_equ = self.style_equ()
        size_equ = self.size_equ()
        pepbroe_equ = self.pepbroe_equ()
        nav_equ = self.nav_equ()

        #rename columns :
        ind_equ=pd.merge(ind_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').\
            rename(columns={'cen_ind_1':'行业集中度'})

        style_equ=pd.merge(style_equ,
                         jjdm_name_map,
                         how='left', on='jjdm')\
            .rename(columns={'cen_lv_rank': '风格集中度'})

        size_equ=pd.merge(size_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').\
            rename(columns={'cen_lv_rank': '规模集中度'})

        pepbroe_equ=pd.merge(pepbroe_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').rename(columns={'PE_rank': 'pe偏离度',
                                    'PB_rank': 'pb偏离度',
                                    'ROE_rank': 'roe偏离度',
                                    '股息率_rank': '股息率偏离度'})

        nav_equ=pd.merge(nav_equ,
                         jjdm_name_map,
                         how='left', on='jjdm').rename(columns={'mean_rank_monthly': '净值偏离度（月）',
                                'std_rank_monthly': '净值偏离度方差（月）',
                                'mean_rank_weekly': '净值偏离度（周）',
                                'std_rank_weekly': '净值偏离度方差（周）',
                                })

        # method_2
        joint_rank = pd.merge(ind_equ, style_equ, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, size_equ, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, pepbroe_equ, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, nav_equ, how='inner', on=['jjdm','jjjc'])

        col_list = ['行业集中度', '风格集中度',
                    '规模集中度', 'pe偏离度',
                    '净值偏离度（月）']

        joint_rank['平均集中度'] = joint_rank[col_list].mean(axis=1)
        joint_rank = joint_rank.sort_values('平均集中度')

        # method_3
        joint_result = joint_rank[(joint_rank[col_list] <= threshield).prod(axis=1) == 1][['jjdm','jjjc','asofdate'] + col_list]

        joint_rank=joint_rank[['jjdm','jjjc','平均集中度','asofdate']+col_list]


        #format change
        ind_equ['行业集中度'] = ind_equ['行业集中度'].map("{:.2%}".format)
        style_equ['风格集中度'] = style_equ['风格集中度'].map("{:.2%}".format)
        size_equ['规模集中度'] = size_equ['规模集中度'].map("{:.2%}".format)

        for col in ['pe偏离度','pb偏离度','roe偏离度','股息率偏离度']:
            pepbroe_equ[col] = pepbroe_equ[col].map("{:.2%}".format)
        for col in ['净值偏离度（月）', '净值偏离度方差（月）', '净值偏离度（周）', '净值偏离度方差（周）']:
            nav_equ[col] = nav_equ[col].map("{:.2%}".format)
        for col in col_list:
            joint_rank[col]=joint_rank[col].map("{:.2%}".format)
            joint_result[col]=joint_result[col].map("{:.2%}".format)
        joint_rank['平均集中度']=joint_rank['平均集中度'].map("{:.2%}".format)

        return ind_equ[0:show_num],style_equ[0:show_num],\
               size_equ[0:show_num],pepbroe_equ[0:show_num],nav_equ[0:show_num],\
               joint_rank[0:show_num],joint_result[0:show_num]

class Leftside:

    @staticmethod
    def stock_left():

        latest_asofdate=\
            pd.read_sql("select max(asofdate) as asofdate from hbs_stock_trading_property ",
                        con=localdb)['asofdate'][0]

        sql="""
        SELECT jjdm,`左侧概率（出持仓前,半年线）_rank`,`左侧概率（出持仓前,年线）_rank`,asofdate 
        from hbs_stock_trading_property where asofdate='{0}'  """\
            .format(latest_asofdate)
        stock_left=pd.read_sql(sql,con=localdb)

        stock_left['stock_left_rank']=stock_left[['左侧概率（出持仓前,半年线）_rank','左侧概率（出持仓前,年线）_rank']]\
            .max(axis=1)

        stock_left=stock_left.sort_values('stock_left_rank',ascending=False)

        return stock_left

    @staticmethod
    def ind_left():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_industry_shift_property_new "
                        , con=localdb)['asofdate'][0]
        sql="""
        SELECT jjdm,项目名,Total_rank,asofdate 
        from hbs_industry_shift_property_new where asofdate='{0}' and (项目名='左侧比率' or 项目名='深度左侧比例') 
        """.format(latest_asofdate)
        ind_left=pd.read_sql(sql,con=localdb)

        ind_left=ind_left.groupby('jjdm').max('Total_rank')
        ind_left.reset_index(inplace=True)
        ind_left['asofdate']=latest_asofdate

        ind_left.sort_values('Total_rank',ascending=False,inplace=True)

        ind_left.rename(columns={'Total_rank':'ind_rank'},inplace=True)

        return ind_left

    @staticmethod
    def value_left():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_shift_property_value "
                        , con=localdb)['asofdate'][0]
        sql="""
        SELECT jjdm,项目名,Total_rank,asofdate 
        from hbs_shift_property_value where asofdate='{0}' and (项目名='左侧比率' or 项目名='深度左侧比例') 
        """.format(latest_asofdate)
        value_left=pd.read_sql(sql,con=localdb)

        value_left=value_left.groupby('jjdm').max('Total_rank')
        value_left.reset_index(inplace=True)
        value_left['asofdate']=latest_asofdate

        value_left.sort_values('Total_rank',ascending=False,inplace=True)
        value_left.rename(columns={'Total_rank': 'value_rank'}, inplace=True)

        return value_left

    @staticmethod
    def size_left():

        latest_asofdate = \
            pd.read_sql("select max(asofdate) as asofdate from hbs_shift_property_size "
                        , con=localdb)['asofdate'][0]
        sql="""
        SELECT jjdm,项目名,Total_rank,asofdate 
        from hbs_shift_property_size where asofdate='{0}' and (项目名='左侧比率' or 项目名='深度左侧比例') 
        """.format(latest_asofdate)
        size_left=pd.read_sql(sql,con=localdb)

        size_left=size_left.groupby('jjdm').max('Total_rank')
        size_left.reset_index(inplace=True)
        size_left['asofdate']=latest_asofdate

        size_left.sort_values('Total_rank',ascending=False,inplace=True)
        size_left.rename(columns={'Total_rank': 'size_rank'}, inplace=True)


        return size_left

    def get_left(self,threshield,show_num=300):

        # method_1
        stock_left = self.stock_left()
        ind_left = self.ind_left()
        value_left = self.value_left()
        size_left = self.size_left()

        #rename columns
        stock_left=stock_left[['jjdm','stock_left_rank','asofdate']]

        stock_left=pd.merge(stock_left,jjdm_name_map,how='left',on='jjdm')\
            .rename(columns={'stock_left_rank':'个股交易左侧率'})
        ind_left = pd.merge(ind_left, jjdm_name_map, how='left', on='jjdm')\
            .rename(columns={'ind_rank':'行业切换左侧率'})
        value_left = pd.merge(value_left, jjdm_name_map, how='left', on='jjdm')\
            .rename(columns={'value_rank':'风格切换左侧率'})
        size_left = pd.merge(size_left, jjdm_name_map, how='left', on='jjdm')\
            .rename(columns={'size_rank':'规模切换左侧率'})


        # method_2
        joint_rank = pd.merge(stock_left, ind_left, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, value_left, how='inner', on=['jjdm','jjjc'])
        joint_rank = pd.merge(joint_rank, size_left, how='inner', on=['jjdm','jjjc'])
        col_list = ['个股交易左侧率', '行业切换左侧率', '风格切换左侧率', '规模切换左侧率']
        joint_rank['平均左侧率'] = joint_rank[col_list].mean(axis=1)

        # method_3
        joint_restult = joint_rank[(joint_rank[col_list] >= (1 - threshield)).prod(axis=1) == 1][['jjdm','jjjc'] + col_list]

        joint_rank = joint_rank.sort_values('平均左侧率', ascending=False)[['jjdm','jjjc', '平均左侧率', 'asofdate_x']+col_list]


        #re format
        stock_left['个股交易左侧率']=stock_left['个股交易左侧率'].map("{:.2%}".format)
        ind_left['行业切换左侧率'] = ind_left['行业切换左侧率'].map("{:.2%}".format)
        value_left['风格切换左侧率'] = value_left['风格切换左侧率'].map("{:.2%}".format)
        size_left['规模切换左侧率'] = size_left['规模切换左侧率'].map("{:.2%}".format)
        for col in col_list:
            joint_rank[col] = joint_rank[col].map("{:.2%}".format)
            joint_restult[col] = joint_restult[col].map("{:.2%}".format)

        joint_rank['平均左侧率']=joint_rank['平均左侧率'].map("{:.2%}".format)

        return stock_left[0:show_num],ind_left[0:show_num],value_left[0:show_num],size_left[0:show_num],joint_rank[0:show_num],joint_restult[0:show_num]

class Size:
    @staticmethod
    def size_property(fre):
        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from hbs_size_property".format(fre), con=localdb)['asofdate'][0]
        sql = "select jjdm, shift_lv, cen_lv, shift_lv_rank, cen_lv_rank,大盘, 中盘, 小盘 ,大盘_rank, 中盘_rank, 小盘_rank, asofdate from hbs_size_property where asofdate='{0}'".format( latest_asofdate)
        size_property = pd.read_sql(sql, con=localdb)
        return size_property

    def get_size(self, fre, show_num=200, shift_ratio_threshold=0.5, centralization_threshold=0.5):
        size_property = self.size_property(fre)
        size_property.columns = ['jjdm', '换手率', '集中度', '换手率排名', '集中度排名','大盘(绝对值）', '中盘(绝对值）', '小盘(绝对值）' ,'大盘', '中盘', '小盘', 'asofdate']
        size_property = size_property.merge(jjdm_name_map, on=['jjdm'], how='left')

        size = size_property[(size_property['换手率排名'] < shift_ratio_threshold) & (size_property['集中度排名'] > centralization_threshold)]
        big_size = size[(size['大盘'] > size['中盘']) & (size['大盘'] > size['小盘'])].sort_values('大盘', ascending=False)
        medium_size = size[(size['中盘'] > size['大盘']) & (size['中盘'] > size['小盘'])].sort_values('中盘', ascending=False)
        small_size = size[(size['小盘'] > size['大盘']) & (size['小盘'] > size['中盘'])].sort_values('小盘', ascending=False)

        big_size = big_size[['jjdm', 'jjjc', 'asofdate', '大盘']]
        medium_size = medium_size[['jjdm', 'jjjc', 'asofdate', '中盘']]
        small_size = small_size[['jjdm', 'jjjc', 'asofdate', '小盘']]
        big_size['大盘'] = big_size['大盘'].map("{:.2%}".format)
        medium_size['中盘'] = medium_size['中盘'].map("{:.2%}".format)
        small_size['小盘'] = small_size['小盘'].map("{:.2%}".format)
        return big_size[0: show_num], medium_size[0: show_num], small_size[0: show_num]

class Value:
    @staticmethod
    def value_property(fre):
        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from hbs_style_property ", con=localdb)['asofdate'][0]
        sql = "select jjdm, shift_lv, cen_lv, shift_lv_rank, cen_lv_rank,成长,价值,成长_rank, 价值_rank, asofdate from hbs_style_property where  asofdate='{0}'".format(latest_asofdate)
        value_property = pd.read_sql(sql, con=localdb)
        return value_property

    @staticmethod
    def holding_property():
        latest_asofdate = pd.read_sql("select max(asofdate) as asofdate from hbs_holding_property", con=localdb)['asofdate'][0]
        sql = "select jjdm, PE_rank, PB_rank, PE_REL_rank, PB_REL_rank, ROE_rank, 股息率_rank, asofdate from hbs_holding_property where asofdate='{0}'".format(latest_asofdate)
        holding_property = pd.read_sql(sql, con=localdb)
        return holding_property

    def get_value(self, fre, show_num=200, shift_ratio_threshold=0.5, centralization_threshold=0.5):
        value_property = self.value_property(fre)
        value_property.columns = ['jjdm', '换手率', '集中度', '换手率排名', '集中度排名','成长（绝对值）', '价值（绝对值）' ,'成长', '价值', 'asofdate']
        value_property = value_property.merge(jjdm_name_map, on=['jjdm'], how='left')

        value = value_property[(value_property['换手率排名'] < shift_ratio_threshold) & (value_property['集中度排名'] > centralization_threshold)]
        growth = value[value['成长'] > value['价值']].sort_values('成长', ascending=False)
        value = value[value['价值'] > value['成长']].sort_values('价值', ascending=False)

        growth = growth[['jjdm', 'jjjc', 'asofdate', '成长']]
        value = value[['jjdm', 'jjjc', 'asofdate', '价值']]
        growth['成长'] = growth['成长'].map("{:.2%}".format)
        value['价值'] = value['价值'].map("{:.2%}".format)

        holding_property = self.holding_property()
        holding_property.columns = ['jjdm', 'PE排名', 'PB排名', 'PE相对行业均值排名', 'PB相对行业均值排名', 'ROE排名', '股息率排名', 'asofdate']
        value_holding_property = value.merge(holding_property.drop('asofdate', axis=1), on=['jjdm'], how='left')
        absolute_pe_value = value_holding_property.sort_values('PE排名')
        absolute_pb_value = value_holding_property.sort_values('PB排名')
        relative_pe_value = value_holding_property.sort_values('PE相对行业均值排名')
        relative_pb_value = value_holding_property.sort_values('PB相对行业均值排名')
        dividend_value = value_holding_property.sort_values('股息率排名', ascending=False)
        reverse_value = value_holding_property.sort_values('ROE排名')
        high_quality_value = value_holding_property.sort_values('ROE排名', ascending=False)

        absolute_pe_value = absolute_pe_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PE排名']]
        absolute_pb_value = absolute_pb_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PB排名']]
        relative_pe_value = relative_pe_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PE相对行业均值排名']]
        relative_pb_value = relative_pb_value[['jjdm', 'jjjc', 'asofdate', '价值', 'PB相对行业均值排名']]
        dividend_value = dividend_value[['jjdm', 'jjjc', 'asofdate', '价值', '股息率排名']]
        reverse_value = reverse_value[['jjdm', 'jjjc', 'asofdate', '价值', 'ROE排名']]
        high_quality_value = high_quality_value[['jjdm', 'jjjc', 'asofdate', '价值', 'ROE排名']]

        absolute_pe_value['PE排名'] = absolute_pe_value['PE排名'].map("{:.2%}".format)
        absolute_pb_value['PB排名'] = absolute_pb_value['PB排名'].map("{:.2%}".format)
        relative_pe_value['PE相对行业均值排名'] = relative_pe_value['PE相对行业均值排名'].map("{:.2%}".format)
        relative_pb_value['PB相对行业均值排名'] = relative_pb_value['PB相对行业均值排名'].map("{:.2%}".format)
        dividend_value['股息率排名'] = dividend_value['股息率排名'].map("{:.2%}".format)
        reverse_value['ROE排名'] = reverse_value['ROE排名'].map("{:.2%}".format)
        high_quality_value['ROE排名'] = high_quality_value['ROE排名'].map("{:.2%}".format)
        return growth[0: show_num], value[0: show_num], absolute_pe_value[0: show_num], absolute_pb_value[0: show_num], \
               relative_pe_value[0: show_num], relative_pb_value[0: show_num], dividend_value[0: show_num], \
               reverse_value[0: show_num], high_quality_value[0: show_num]

class Report:


    @staticmethod
    def get_silly_summary(doc,equ_list,equ_name_list,fold,
                          paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num):

        from PIL import Image
        from docx.shared import Inches, Pt


        def style_analysis(df,colname):
            style_style=''
            for style in df.index:
                style_style+=str(style)+"占比"+str(round(df.loc[style][colname],2))+"%,"

            return  style_style

        pic_name_list = ['池列表.png',
                         '行业风格分布.png', '1级行业持仓分布_占持仓比例.png',
                         '1级行业持仓分布_占持仓比例排名.png',
                         '2级行业持仓分布_占持仓比例.png',
                         '2级行业持仓分布_占持仓比例排名.png',
                         '3级行业持仓分布_占持仓比例.png',
                         '3级行业持仓分布_占持仓比例排名.png',
                         '风格类型分布.png',
                         '风格专注型分布.png',
                         '成长价值持仓占比.png',
                         '成长价值持仓占比排名.png',
                         '规模风格类型分布.png',
                         '专注型风格分布.png',
                         '个股持仓特征.png',
                         '大中小盘持仓占比.png',
                         '大中小盘持仓占比排名.png',
                         '个股风格类型分布A.png',
                         '个股风格类型分布B.png',
                         '左侧类型分布.png',
                         '新股偏好分布.png',
                         ]

        for i in range(len(equ_list)):
            equ=equ_list[i]
            name=equ_name_list[i]

            doc.add_paragraph("")
            paragraphs_count += 1

            p=doc.add_paragraph("{2}.{3} {0}：对{1}排名前30的基金进行了统计画像，统计结果如下：".format(name, name,title_num,i+1))
            p.style.font.size=Pt(10)
            p.paragraph_format.first_line_indent=p.style.font.size*2
            p.paragraph_format.line_spacing = Pt(20)  # 行距
            string_paragraphs_list.append(paragraphs_count)
            paragraphs_count+=1


            industry_style, class_list, style_type_dis, style_incline_dis, \
            style_weight_dis, style_type_dis2, style_incline_dis2, style_weight_dis2, \
            stock_style_a, stock_style_b, stock_left, stock_new, stock_fin = pla.pool_picturing(equ, returndata=True,save_local_file=True)

            for pic in pic_name_list:
                file_path=fold+"\\"+pic
                try:
                    doc.add_picture(file_path, width=Inches(3), height=Inches(2.8))
                    pic_paragraphs_list.append(paragraphs_count)
                    paragraphs_count += 1
                except Exception as e:
                    pic_temp = Image.open(file_path)
                    pic_temp.save(pic_temp)
                    doc.add_picture(file_path, width=Inches(4.5), height=Inches(2.5))

            ind_style=style_analysis(industry_style,'num')
            style_style=style_analysis(style_type_dis,'风格类型')
            style_weight=style_analysis(style_weight_dis,'成长价值权重分布')
            size_style=style_analysis(style_type_dis2,'规模风格类型')
            size_weight=style_analysis(style_weight_dis2,'成长价值权重分布')
            left_property=style_analysis(stock_left,'左侧类型分布')
            new_stock_property=style_analysis(stock_new,'新股偏好分布')
            stock_hld_property=style_analysis(stock_fin,'个股持仓特征')

            summary="总结来看：从行业上看，{0}池中".format(name)+ind_style+\
                    "一级行业中，{0}占比较高".format(util.list_sql_condition(class_list[0].index[0:3].tolist()))+\
                    "从风格上看，{0}池中".format(name)+style_style+"具体的持仓来看"+style_weight+"。"+\
                    "从规模上看，{0}池中".format(name) + size_style + "具体的持仓来看" + size_weight + "。"+\
                    "从左侧属性上看，{0}池中".format(name) + left_property +"。"+\
                    "从新股属性上看，{0}池中".format(name) + new_stock_property + "。" +\
                    "从个股持仓属性上看，{0}池中".format(name) + stock_hld_property

            doc.add_paragraph("")
            paragraphs_count += 1

            p=doc.add_paragraph(summary)
            p.style.font.size = Pt(10)
            p.paragraph_format.first_line_indent=p.style.font.size*2
            p.paragraph_format.line_spacing = Pt(20)  # 行距
            string_paragraphs_list.append(paragraphs_count)
            paragraphs_count+=1

        return doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list

    def filter_pool_report(self):

        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        #save png 2 doc

        equclass=Equilibrium()
        ind_equ, style_equ, size_equ,pepbroe_equ, nav_equ,\
        joint_rank,joint_result=equclass.get_equilibrium(threshield=0.3,show_num=30)
        equ_list=[joint_rank,joint_result]
        equ_name_list=['均衡平均','均衡交集']

        leftclass=Leftside()
        stock_left, ind_left, value_left,size_left,\
        joint_rank, joint_restult=leftclass.get_left(threshield=0.4,show_num=30)
        left_list=[joint_rank, joint_restult]
        left_name_list = ['左侧平均', '左侧交集']

        sizeclass = Size()
        size_list=sizeclass.get_size(fre='M', show_num=30,
                                                               shift_ratio_threshold=0.5,
                                                               centralization_threshold=0.5)
        size_name_list = ['大盘', '中盘', '小盘']

        valueclass = Value()
        value_list=valueclass.get_value(fre='M', show_num=30,
                                        shift_ratio_threshold=0.5, centralization_threshold=0.5)
        value_list=value_list[0:2]
        value_name_list = ['成长', '价值']

        doc = Document()
        fold=r"E:\GitFolder\hbshare\fe\mutual_analysis"

        pic_paragraphs_list=[]
        string_paragraphs_list=[]
        paragraphs_count=0

        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc,equ_list,equ_name_list,fold
                               ,paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=2)
        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc, left_list, left_name_list, fold,
                               paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=3)
        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc, size_list, size_name_list, fold,
                               paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=4)
        doc,paragraphs_count,pic_paragraphs_list,string_paragraphs_list=self.get_silly_summary(doc, value_list, value_name_list, fold,
                               paragraphs_count,pic_paragraphs_list,string_paragraphs_list,title_num=5)

        for j in pic_paragraphs_list:
            doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for j in string_paragraphs_list:
            doc.paragraphs[j].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


        doc.save(r"E:\GitFolder\hbshare\fe\mutual_analysis\筛选报告\筛选池画像报告.docx")


        print('')

class Advance_man:

    @staticmethod
    def advance_man_from3ind():
        target_ind=['电池化学品','锂电池','锂电专用设备','燃料电池','蓄电池及其他电池','光伏电池组件',
                    '硅料硅片','光伏辅材','光伏加工设备','逆变器','风电零部件','风电整机','电网自动化设备',
                    '电工仪器仪表','配电设备','输变电设备','线缆部件及其他','综合电力设备商','工控设备',
                    '机器人','激光设备','其他自动化设备','工程机械器件','工程机械整机','机床工具','金属制品',
                    '磨具磨料','仪器仪表','制冷空调设备','其他通用设备','纺织服装设备','楼宇设备','能源及重型设备',
                    '其他专用设备','轨交设备Ⅲ','半导体材料','半导体设备','分立器件','集成电路封测','集成电路制造',
                    '模拟芯片设计','数字芯片设计','印制电路板','被动元件','LED','光学元件','面板','品牌消费电子',
                    '消费电子零部件及组装','电子化学品Ⅲ','其他电子Ⅲ','电动乘用车','综合乘用车','商用载货车',
                    '商用载客车','底盘与发动机系统','汽车电子电气系统','轮胎轮毂','车身附件及饰件','其他汽车零部件',
                    '航空装备Ⅲ','航天装备Ⅲ','军工电子Ⅲ','地面兵装Ⅲ','航海装备Ⅲ','医疗设备','医疗耗材','体外诊断',
                    '疫苗','血液制品','其他生物制品','化学制剂','原料药','医疗研发外包','诊断服务','医美耗材','冰洗',
                    '空调','（厨卫电器）厨房电器','（黑色家电）彩电','（小家电）清洁小家电','通信网络设备及器件',
                    '通信线缆及配套','通信终端及配件','其他通信设备','安防设备','其他计算机设备','光伏发电','大气治理',
                    '固废治理','（环保设备）环保设备Ⅲ','油气及炼化工程','油田服务']
        sql="SELECT jjdm,sum(`占持仓比例(时序均值)`) as adv_man from jjpic_industry_detail_3 where `行业名称` in({0}) GROUP BY jjdm "\
            .format(util.list_sql_condition(target_ind))
        adv_man=pd.read_sql(sql,con=localdb).sort_values('adv_man',ascending=False)

        return  adv_man
    @staticmethod
    def advance_man_wording():

        key_words=['先进制造','高端制造','智能制造','国产替代','专精特新','新能源','5G','芯片','半导体']
        jjname=pd.DataFrame(data=util.get_mutual_stock_funds('20211231'),columns=['jjdm'])

        sql="select jjdm,jjmc from st_fund.t_st_gm_jjxx "
        jjname=pd.merge(jjname,hbdb.db2df(sql,db='funduser'),how='left',on='jjdm')

        orconditon=[False]*len(jjname)
        for word in key_words:
            orconditon=(orconditon)|(jjname['jjmc'].str.contains(word))

        return  jjname.loc[orconditon]
    @staticmethod
    def advance_man_fromticker(jjdm_list):
        from hbshare.fe.mutual_analysis import holdind_based as hb

        hld=hb.read_hld_fromdb(jjdm_list,start_date='20211231',end_date='20211231')

        ticker_list=pd.read_excel(r"E:\GitFolder\hbshare\fe\mutual_analysis\指数成份汇总.xlsx",sheet_name='成份汇总')
        ticker_list['Wind代码']=ticker_list['Wind代码'].astype(str).str.replace("'", "")
        hld=pd.merge(hld,ticker_list,how='left',left_on='zqdm',right_on='Wind代码')
        hld=hld[hld['Wind代码'].notnull()]
        hld=hld.groupby('jjdm').sum()
        hld['zjbl']=hld['zjbl'].rank(method='min')/len(hld)
        hld.rename(columns={'zjbl':'zjbl_ticker'})
        return hld.reset_index()
    @staticmethod
    def advanve_ols(jjdm_list):
        from hbshare.fe.mutual_analysis import  nav_based as nb
        import statsmodels.api as sm

        index_ret=pd.read_excel(r"E:\GitFolder\hbshare\fe\mutual_analysis\指数成份汇总.xlsx",
                                  sheet_name='走势').drop('881001.WI',axis=1)
        index_ret[['930820.CSI','931167.CSI','866003.WI']]=\
            index_ret[['930820.CSI','931167.CSI','866003.WI']].pct_change()
        index_ret['Date']=index_ret['Date'].astype(str).str.replace("-","")
        index_ret=index_ret.iloc[2:]
        index_ret.set_index('Date',inplace=True)
        index_ret = index_ret.sort_index().loc['20190311':]

        result=[[],[],[]]
        para_df=pd.DataFrame(data=jjdm_list,columns=['jjdm'])

        for jjdm in jjdm_list:
            print(jjdm)
            jj_ret = nb.get_jj_daily_ret([jjdm])
            jj_ret.index=jj_ret.index.astype(str)
            olsdf = pd.merge(index_ret,jj_ret, how='left', left_index=True, right_index=True).fillna(0)

            laster_3_year = (datetime.datetime.strptime(str(olsdf.index[-1]),
                                                                                          '%Y%m%d') - datetime.timedelta(
                days=365*3)).strftime('%Y%m%d')

            for j in range(3):
                ind=['930820.CSI','931167.CSI','866003.WI'][j]
                para = sm.OLS(olsdf.loc[laster_3_year:][jjdm].values,
                              olsdf.loc[laster_3_year:][
                                  ind].values).fit().params.tolist()
                result[j].append(para[0])

        for j in range(3):
            ind = ['930820.CSI', '931167.CSI', '866003.WI'][j]+'_rank'
            para_df[ind]=result[j]
            para_df[ind]=para_df[ind].rank(method='min')/len(para_df)

        return  para_df

    def get_advance_man(self):

        jjdm_list = util.get_mutual_stock_funds('20211231')

        adv_man_3ind = self.advance_man_from3ind()
        adv_man_ols=self.advanve_ols(jjdm_list)
        adv_man_ticker=self.advance_man_fromticker(jjdm_list)
        adv_man_words=self.advance_man_wording()

        adv_man=pd.merge(adv_man_3ind,adv_man_words
                         ,how='left',on='jjdm')
        adv_man=pd.merge(adv_man,adv_man_ticker
                         ,how='left',on='jjdm')
        adv_man=pd.merge(adv_man,adv_man_ols
                         ,how='left',on='jjdm')

        adv_man.to_excel('先进制造_基于三级行业列表.xlsx')

        print('Done')

if __name__ == '__main__':

    am=Advance_man()
    am.get_advance_man()

    # rp=Report()
    # rp.filter_pool_report()

    #advance_man()


    #
    # equclass=Equilibrium()
    # ind_equ, style_equ,size_equ, \
    # pepbroe_equ, nav_equ,joint_rank, \
    # joint_result=equclass.get_equilibrium(threshield=0.3,show_num=100)
    #
    #
    #
    # plot=functionality.Plot(800,800)
    #
    # plot.plotly_table(ind_equ, 800, 'asdf')
    # pla.pool_picturing(ind_equ)
    #
    # plot.plotly_table(style_equ, 800, 'asdf')
    # pla.pool_picturing(style_equ)
    #
    # plot.plotly_table(size_equ, 800, 'asdf')
    # pla.pool_picturing(size_equ)
    #
    # plot.plotly_table(pepbroe_equ, 800, 'asdf')
    # pla.pool_picturing(pepbroe_equ)
    #
    # plot.plotly_table(nav_equ, 800, 'asdf')
    # pla.pool_picturing(nav_equ)
    #
    # plot.plotly_table(joint_result,800,'asdf')
    # pla.pool_picturing(joint_result)
    #
    # plot.plotly_table(joint_rank, 800, 'asdf')
    # pla.pool_picturing(joint_rank)

    #

    # leftclass=Leftside()
    #
    # stock_left, ind_left, value_left, \
    # size_left, joint_rank, joint_restult=leftclass.get_left(threshield=0.4,show_num=100)
    #
    #
    # plot=functionality.Plot(800,800)
    #
    # plot.plotly_table(stock_left,800,'asdf')
    # pla.pool_picturing(stock_left)
    #
    # plot.plotly_table(ind_left, 800, 'asdf')
    # pla.pool_picturing(ind_left)
    #
    # plot.plotly_table(value_left, 800, 'asdf')
    # pla.pool_picturing(value_left)
    #
    # plot.plotly_table(size_left, 800, 'asdf')
    # pla.pool_picturing(size_left)
    #
    # plot.plotly_table(joint_rank, 800, 'asdf')
    # pla.pool_picturing(joint_rank)

    # plot.plotly_table(joint_restult, 800, 'asdf')
    # pla.pool_picturing(joint_restult)

    # sizeclass = Size()
    # big_size, medium_size, small_size = sizeclass.get_size(fre='M', show_num=100,
    #                                                        shift_ratio_threshold=0.5,
    #                                                        centralization_threshold=0.5)

    #plot = functionality.Plot(800, 800)
    # plot.plotly_table(big_size, 800, 'asdf')
    # pla.pool_picturing(big_size)
    #
    # plot.plotly_table(medium_size, 800, 'asdf')
    # pla.pool_picturing(medium_size)
    #
    # plot.plotly_table(small_size, 800, 'asdf')
    # pla.pool_picturing(small_size)


    # valueclass = Value()
    # growth, value, absolute_pe_value, absolute_pb_value, relative_pe_value, relative_pb_value, \
    # dividend_value, reverse_value, high_quality_value = valueclass.get_value(fre='M', show_num=200, shift_ratio_threshold=0.5, centralization_threshold=0.5)
    #
    # plot = functionality.Plot(800, 800)
    # plot.plotly_table(growth, 800, 'asdf')
    # pla.pool_picturing(growth)
    #
    # plot.plotly_table(value, 800, 'asdf')
    # pla.pool_picturing(value)
    #
    # plot.plotly_table(absolute_pe_value, 800, 'asdf')
    # pla.pool_picturing(absolute_pe_value)
    #
    # plot.plotly_table(absolute_pb_value, 800, 'asdf')
    # pla.pool_picturing(absolute_pb_value)
    #
    # plot.plotly_table(relative_pe_value, 800, 'asdf')
    # pla.pool_picturing(relative_pe_value)
    #
    # plot.plotly_table(relative_pb_value, 800, 'asdf')
    # pla.pool_picturing(relative_pb_value)
    #
    # plot.plotly_table(dividend_value, 800, 'asdf')
    # pla.pool_picturing(dividend_value)
    #
    # plot.plotly_table(reverse_value, 800, 'asdf')
    # pla.pool_picturing(reverse_value)
    #
    # plot.plotly_table(high_quality_value, 800, 'asdf')
    # pla.pool_picturing(high_quality_value)
    #
    # print("Done")


